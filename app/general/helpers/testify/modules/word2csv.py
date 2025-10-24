from docx import Document    # use main import (better)
import os
from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree as ET

def _extract_tables_from_raw_docx(path):
    """Try multiple decodings of word/document.xml and extract tables as
    list-of-rows-of-cells. Returns [] on failure."""
    encodings_to_try = ("utf-8", "cp1251", "koi8-r", "iso8859-5")
    with ZipFile(path) as z:
        xml_bytes = z.read('word/document.xml')
    for enc in encodings_to_try:
        try:
            xml_text = xml_bytes.decode(enc, errors='replace')
            root = ET.fromstring(xml_text)
        except Exception:
            continue

        tables = []
        # find all <w:tbl> elements anywhere in document
        for tbl in root.iter():
            if tbl.tag.endswith('}tbl'):
                rows = []
                for tr in tbl.iter():
                    if tr.tag.endswith('}tr'):
                        row_cells = []
                        for tc in tr.iter():
                            if tc.tag.endswith('}tc'):
                                # gather all <w:t> inside the cell
                                parts = []
                                for t in tc.iter():
                                    if t.tag.endswith('}t') and t.text:
                                        parts.append(t.text)
                                cell_text = ''.join(parts).replace('\r', '').replace('\n', ' ').strip()
                                row_cells.append(cell_text)
                        if row_cells:
                            rows.append(row_cells)
                if rows:
                    tables.append(rows)
        if tables:
            return tables
    return []

def _get_table_rows(filepath, fallback_if_garbled=True):
    """Return first table as list of rows (each row is list of cell texts).
       Try python-docx first, then fallback to raw xml extraction if needed."""
    # try python-docx
    try:
        doc = Document(filepath)
        if doc.tables:
            table = doc.tables[0]
            rows = []
            for r in table.rows:
                cells = [c.text.replace('\r', '').replace('\n', ' ').strip() for c in r.cells]
                rows.append(cells)
            # detect garbled: either all empty or contains replacement char
            if not rows or all(all(not cell for cell in row) for row in rows) or any('\ufffd' in cell for row in rows for cell in row):
                # fallback to raw xml extraction
                raw_tables = _extract_tables_from_raw_docx(filepath)
                if raw_tables:
                    return raw_tables[0]
            else:
                return rows
    except Exception:
        pass

    # final fallback: raw extraction only
    raw_tables = _extract_tables_from_raw_docx(filepath)
    return raw_tables[0] if raw_tables else []

def word2csv(doc_path, input_file, test_id):
    try:
        # variables 
        doc_path = Path('./public/') / doc_path
        data_path = doc_path / 'csv'
        file_name = input_file.split('.')[0]
        output_file = f'{file_name}.csv'

        # create output folder if do not exists
        os.makedirs(data_path, exist_ok=True)

        # --- get table rows (each row is list of cell texts) ---
        docx_file_path = doc_path / input_file
        rows = _get_table_rows(str(docx_file_path))

        if not rows:
            print("No table found or could not extract text.")
            return None

        data = []
        # manual header
        keys = ("Question", "A", "B", "C", "D")

        for i, row_cells in enumerate(rows):
            # if row has more cells than keys, keep only first len(keys)
            # if fewer, zip will truncate - acceptable behavior
            row_data = dict(zip(keys, row_cells))
            data.append(row_data)

        # Specify UTF-8 encoding here to avoid encoding issues
        with open(data_path / output_file, 'w', encoding='utf-8') as file:
            csv_header = '"test_id","question","option_a","option_b","option_c","option_d"'
            file.write(csv_header)
            file.write('\n')
            for row in data:
                question = row.get("Question", "")
                variation1 = row.get("A", "")
                variation2 = row.get("B", "")
                variation3 = row.get("C", "")
                variation4 = row.get("D", "")

                # Basic CSV quoting (keeps your original style)
                row_text = (
                    f'"{test_id}","{question}","{variation1}","{variation2}","{variation3}","{variation4}"'
                )
                file.write(row_text)
                file.write('\n')

        print("Word has been converted to csv successfully.")
        return output_file
    except Exception as e:
        print(e)
        return None
